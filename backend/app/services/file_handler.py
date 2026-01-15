"""
File handler service for parsing and exporting documents.
Supports PDF, DOCX, and TXT formats with highlighting.
"""
import base64
import io
import logging
from typing import List, BinaryIO, Optional
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
import pdfplumber
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

logger = logging.getLogger(__name__)


class FileHandler:
    """
    Service for handling file operations: parsing and exporting.
    
    Supports:
    - Parsing: PDF, DOCX, TXT
    - Exporting: DOCX, PDF with color highlighting
    """
    
    def __init__(self):
        """Initialize the file handler."""
        pass
    
    def parse_file(self, file_content: bytes, filename: str) -> str:
        """
        Parse a file and extract text.
        
        Args:
            file_content: Raw file content
            filename: Original filename (used to determine type)
            
        Returns:
            Extracted text
        """
        extension = filename.lower().split('.')[-1] if '.' in filename else ''
        
        try:
            if extension == 'pdf':
                return self._parse_pdf(file_content)
            elif extension in ['docx', 'doc']:
                return self._parse_docx(file_content)
            elif extension == 'txt':
                return self._parse_txt(file_content)
            else:
                # Try to parse as text
                return self._parse_txt(file_content)
        except Exception as e:
            logger.error(f"Error parsing file {filename}: {e}")
            raise ValueError(f"Failed to parse file: {e}")
    
    def _parse_pdf(self, content: bytes) -> str:
        """Parse PDF file."""
        text_parts = []
        
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        return '\n'.join(text_parts)
    
    def _parse_docx(self, content: bytes) -> str:
        """Parse DOCX file."""
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        return '\n'.join(text_parts)
    
    def _parse_txt(self, content: bytes) -> str:
        """Parse text file."""
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    return content.decode(encoding)
                except:
                    continue
            raise ValueError("Unable to decode text file")
    
    def export_docx(self, sentences: List[dict], original_filename: Optional[str] = None) -> bytes:
        """
        Export analyzed sentences to DOCX with color highlighting.
        
        Args:
            sentences: List of sentence analysis results
            original_filename: Original filename (unused, kept for API compatibility)
            
        Returns:
            DOCX file as bytes
        """
        doc = Document()
        
        # Add title
        title = doc.add_heading('AI Text Detection Analysis', 0)
        
        # Add each sentence with highlighting
        for sentence_data in sentences:
            text = sentence_data['text']
            classification = sentence_data['classification']
            confidence = sentence_data.get('confidence', 0.0)
            
            # Create paragraph
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            
            # Apply highlighting based on classification
            if classification == 'human':
                # Green background
                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            elif classification == 'suspicious':
                # Yellow background
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif classification == 'ai':
                # Red background
                run.font.highlight_color = WD_COLOR_INDEX.RED
            
            # Add confidence as small text
            confidence_text = paragraph.add_run(f" [{confidence:.2%}]")
            confidence_text.font.size = Pt(8)
            confidence_text.font.color.rgb = RGBColor(128, 128, 128)
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    def export_pdf(self, sentences: List[dict], original_filename: Optional[str] = None) -> bytes:
        """
        Export analyzed sentences to PDF with color highlighting.
        
        Args:
            sentences: List of sentence analysis results
            original_filename: Original filename (unused, kept for API compatibility)
            
        Returns:
            PDF file as bytes
        """
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = styles['Heading1']
        
        # Add title
        title = Paragraph("AI Text Detection Analysis", title_style)
        elements.append(title)
        elements.append(Spacer(1, 0.2 * inch))
        
        # Create custom styles for highlighting
        normal_style = styles['Normal']
        
        # Add each sentence with background color
        for sentence_data in sentences:
            text = sentence_data['text']
            classification = sentence_data['classification']
            confidence = sentence_data.get('confidence', 0.0)
            
            # Determine background color
            if classification == 'human':
                bg_color = HexColor('#90EE90')  # Light green
            elif classification == 'suspicious':
                bg_color = HexColor('#FFFFE0')  # Light yellow
            elif classification == 'ai':
                bg_color = HexColor('#FFB6C1')  # Light red
            else:
                bg_color = None
            
            # Create custom style with background
            custom_style = ParagraphStyle(
                'CustomStyle',
                parent=normal_style,
                backColor=bg_color,
                spaceBefore=6,
                spaceAfter=6
            )
            
            # Create paragraph with confidence
            full_text = f"{text} <font size=8 color='gray'>[{confidence:.2%}]</font>"
            para = Paragraph(full_text, custom_style)
            elements.append(para)
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    def decode_base64_file(self, base64_content: str) -> bytes:
        """
        Decode base64 encoded file content.
        
        Args:
            base64_content: Base64 encoded string
            
        Returns:
            Decoded bytes
        """
        # Remove data URL prefix if present
        if ',' in base64_content:
            base64_content = base64_content.split(',', 1)[1]
        
        return base64.b64decode(base64_content)
